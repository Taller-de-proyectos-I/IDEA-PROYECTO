import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─────────────────────────────────────────
# PAGE SETUP – A4, narrow margins
# ─────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2.5)

# ─────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────
def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_paragraph(doc, text="", alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  font_name="Times New Roman", font_size=12,
                  bold=False, italic=False, color=None,
                  space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_font(run, font_name, font_size, bold, italic, color)
    return p

def add_heading_custom(doc, text, level=1):
    colors = {
        1: (0, 70, 127),   # dark blue
        2: (0, 100, 160),
        3: (40, 40, 40),
    }
    sizes  = {1: 14, 2: 13, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.font.name  = "Times New Roman"
    run.font.size  = Pt(sizes[level])
    run.font.bold  = True
    run.font.color.rgb = RGBColor(*colors[level])
    return p

def shade_cell(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        if edge in kwargs:
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'),   kwargs[edge].get('val','single'))
            tag.set(qn('w:sz'),    kwargs[edge].get('sz','4'))
            tag.set(qn('w:space'),'0')
            tag.set(qn('w:color'), kwargs[edge].get('color','auto'))
            tcBorders.append(tag)
    tcPr.append(tcBorders)

def cell_text(cell, text, bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

# ─────────────────────────────────────────
# PAGE BREAK HELPER
# ─────────────────────────────────────────
def add_page_break(doc):
    doc.add_page_break()

# ═══════════════════════════════════════════════════
#  CARATULA
# ═══════════════════════════════════════════════════
# Universidad
p = add_paragraph(doc,
    "UNIVERSIDAD NACIONAL DE SAN ANTONIO ABAD DEL CUSCO",
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    font_size=13, bold=True, color=(0,70,127),
    space_before=20, space_after=4)

add_paragraph(doc,
    "Escuela Profesional de Ingeniería Informática y de Sistemas",
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    font_size=12, italic=True, color=(40,40,40),
    space_before=0, space_after=20)

# Línea decorativa
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(20)
run = p.add_run("─" * 65)
run.font.color.rgb = RGBColor(0,70,127)
run.font.size = Pt(11)

# Título del documento
add_paragraph(doc,
    "EVALUACIÓN HEURÍSTICA DE NIELSEN",
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    font_size=18, bold=True, color=(0,70,127),
    space_before=10, space_after=10)

add_paragraph(doc,
    "Aplicada al Sistema Web PlantAndes",
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    font_size=14, bold=True, italic=True, color=(0,100,160),
    space_before=0, space_after=30)

# Nombre del sistema / proyecto
add_paragraph(doc,
    "\"Implementación de sistema de visión por computadora con CNN\npara diagnóstico de enfermedades en hojas de cultivos andinos\nen comunidades de Cusco\"",
    alignment=WD_ALIGN_PARAGRAPH.CENTER,
    font_size=13, bold=True, color=(30,30,30),
    space_before=0, space_after=30)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(30)
run2 = p2.add_run("─" * 65)
run2.font.color.rgb = RGBColor(0,70,127)
run2.font.size = Pt(11)

# Integrantes
add_paragraph(doc, "INTEGRANTES:", alignment=WD_ALIGN_PARAGRAPH.CENTER,
              font_size=12, bold=True, color=(0,70,127), space_before=10, space_after=6)

members = [
    "Yersson Calderon Romero",
    "Alejandro Humberto Guevara Valdivia",
]
for m in members:
    add_paragraph(doc, m, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                  font_size=12, space_before=2, space_after=4)

add_paragraph(doc, "", space_before=10, space_after=4)

# Evaluador / Docente
add_paragraph(doc, "EVALUADOR:", alignment=WD_ALIGN_PARAGRAPH.CENTER,
              font_size=12, bold=True, color=(0,70,127), space_before=10, space_after=6)
add_paragraph(doc, "Oscar Añazco Durand",
              alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=12, space_before=2, space_after=20)

# Asignatura y fecha
add_paragraph(doc, "Asignatura: Taller de Proyectos I",
              alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=12, italic=True,
              space_before=10, space_after=4)
add_paragraph(doc, "Cusco – Perú, 2026",
              alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=12, bold=True,
              space_before=4, space_after=0)

add_page_break(doc)

# ═══════════════════════════════════════════════════
#  INTRODUCCIÓN
# ═══════════════════════════════════════════════════
add_heading_custom(doc, "I. INTRODUCCIÓN", 1)

intro_text = (
    "La evaluación heurística es un método de inspección de usabilidad propuesto por Jakob Nielsen (1994) "
    "que permite identificar problemas de usabilidad en una interfaz de usuario sin necesidad de contar con "
    "usuarios reales. En este método, un conjunto de evaluadores examina la interfaz y juzga su conformidad "
    "con los diez principios reconocidos de usabilidad denominados 'Heurísticas de Nielsen'.\n\n"
    "El presente documento aplica esta metodología al sistema web PlantAndes, una aplicación desarrollada "
    "con el framework Flask (Python) que implementa redes neuronales convolucionales (CNN) para el "
    "diagnóstico automatizado de enfermedades en hojas de cultivos andinos. La plataforma está orientada "
    "a agricultores y técnicos agrícolas de la región Cusco, por lo que la usabilidad y accesibilidad "
    "de su interfaz son factores críticos para su adopción exitosa."
)
add_paragraph(doc, intro_text, font_size=12, space_after=8)

# ═══════════════════════════════════════════════════
#  OBJETIVO
# ═══════════════════════════════════════════════════
add_heading_custom(doc, "II. OBJETIVO", 1)
add_paragraph(doc,
    "Identificar y documentar los problemas de usabilidad presentes en la interfaz web del sistema "
    "PlantAndes mediante la aplicación de las 10 heurísticas de Jakob Nielsen, con el fin de proponer "
    "mejoras que optimicen la experiencia del usuario agricultor y técnico agrícola.",
    font_size=12, space_after=8)

# ═══════════════════════════════════════════════════
#  DESCRIPCIÓN DEL SISTEMA
# ═══════════════════════════════════════════════════
add_heading_custom(doc, "III. DESCRIPCIÓN DEL SISTEMA EVALUADO", 1)

desc_items = [
    ("Nombre del sistema:", "PlantAndes"),
    ("Tipo de sistema:", "Aplicación web monolítica"),
    ("Tecnologías:", "Flask (Python), HTML5, CSS3, JavaScript, SQLite, PyTorch (CNN), DeepSeek API"),
    ("Usuarios objetivo:", "Agricultores y técnicos agrícolas de la región Cusco, Perú"),
    ("Funcionalidades principales:",
     "Diagnóstico de enfermedades en hojas mediante IA, asistente conversacional AndesGPT, "
     "historial de diagnósticos, panel de administración, soporte multilenguaje (Español / Quechua)"),
    ("URL del sistema:", "http://127.0.0.1:5000 (entorno local de desarrollo)"),
]

for label, value in desc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    r1 = p.add_run(f"{label} ")
    set_font(r1, bold=True, size=12)
    r2 = p.add_run(value)
    set_font(r2, size=12)

add_paragraph(doc, "", space_after=6)

# ═══════════════════════════════════════════════════
#  METODOLOGÍA
# ═══════════════════════════════════════════════════
add_heading_custom(doc, "IV. METODOLOGÍA", 1)

met_text = (
    "La evaluación se realiza siguiendo la metodología de Jakob Nielsen, utilizando la escala de "
    "severidad estándar para cada problema encontrado:\n"
)
add_paragraph(doc, met_text, font_size=12, space_after=4)

severity_items = [
    ("0", "No es un problema de usabilidad."),
    ("1", "Cosmético — Solo necesita ser reparado si hay tiempo disponible."),
    ("2", "Menor — Baja prioridad de reparación."),
    ("3", "Mayor — Alta prioridad, debe ser reparado."),
    ("4", "Catastrófico — Imperativo repararlo antes del lanzamiento."),
]
for num, desc in severity_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Cm(1)
    p.paragraph_format.space_before  = Pt(2)
    p.paragraph_format.space_after   = Pt(2)
    r1 = p.add_run(f"  {num} – ")
    set_font(r1, bold=True, size=12)
    r2 = p.add_run(desc)
    set_font(r2, size=12)

add_paragraph(doc, "", space_after=6)

# ═══════════════════════════════════════════════════
#  EVALUADORES
# ═══════════════════════════════════════════════════
add_heading_custom(doc, "V. EVALUADORES", 1)

evaluators_table = [
    ("Evaluador",          "Rol en el Proyecto",           "Perfil"),
    ("Yersson Calderon Romero",     "Líder / ML Engineer",         "Desarrollador Backend, entrenamiento CNN"),
    ("Alejandro Guevara Valdivia",  "Coordinador / Frontend Dev",  "Desarrollo de interfaz web, UX"),
    ("Oscar Añazco Durand",         "Tester / Evaluador externo",  "Control de calidad y validación de entregables"),
]

tbl = doc.add_table(rows=len(evaluators_table), cols=3)
tbl.style = 'Table Grid'
col_widths = [Cm(5.5), Cm(5.5), Cm(6)]

header_colors = ["004680", "004680", "004680"]
for i, row_data in enumerate(evaluators_table):
    row = tbl.rows[i]
    for j, (cell_data, width) in enumerate(zip(row_data, col_widths)):
        cell = row.cells[j]
        cell.width = width
        is_header = (i == 0)
        shade_cell(cell, header_colors[j] if is_header else ("E8F0F8" if i % 2 == 0 else "FFFFFF"))
        cell_text(cell, cell_data,
                  bold=is_header,
                  size=11,
                  color=((255,255,255) if is_header else (30,30,30)),
                  align=WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT)

add_paragraph(doc, "", space_after=8)

add_page_break(doc)

# ═══════════════════════════════════════════════════
#  HEURÍSTICAS – TABLA PRINCIPAL
# ═══════════════════════════════════════════════════
add_heading_custom(doc, "VI. RESULTADOS DE LA EVALUACIÓN HEURÍSTICA", 1)

add_paragraph(doc,
    "A continuación se presentan los resultados de la evaluación de cada una de las 10 heurísticas "
    "de Nielsen aplicadas a la interfaz del sistema PlantAndes. Para cada heurística se identifican "
    "los problemas encontrados, su nivel de severidad y las recomendaciones de mejora.",
    font_size=12, space_after=10)

# ─────────────────────────────────────────────────────────────────────────────
# DATA: (Heurística, Descripción_breve, [Problema, Área, Severidad, Recomendación])
# ─────────────────────────────────────────────────────────────────────────────
heuristicas = [
    {
        "num":   "H1",
        "title": "Visibilidad del Estado del Sistema",
        "desc":  "El sistema debe mantener informados a los usuarios sobre lo que está ocurriendo, "
                 "mediante retroalimentación apropiada dentro de un tiempo razonable.",
        "problemas": [
            {
                "problema":        "Al subir una imagen para diagnóstico, no se muestra un indicador "
                                   "de progreso mientras el modelo CNN procesa la imagen. El usuario no sabe "
                                   "si el sistema está trabajando o se ha colgado.",
                "area":            "Motor de Diagnóstico IA – Página /index",
                "severidad":       "3 – Mayor",
                "recomendacion":   "Implementar un spinner o barra de progreso animada que sea visible "
                                   "durante todo el proceso de análisis. Mostrar un mensaje como "
                                   "'Analizando imagen...' para tranquilizar al usuario.",
            },
            {
                "problema":        "El sistema no notifica visualmente al usuario cuando se ha guardado "
                                   "un diagnóstico en la base de datos exitosamente, salvo por mensajes "
                                   "flash que desaparecen rápidamente.",
                "area":            "Historial de Diagnósticos",
                "severidad":       "2 – Menor",
                "recomendacion":   "Añadir notificaciones persistentes tipo 'toast' con confirmación "
                                   "de guardado y opción para ver el historial inmediatamente.",
            },
        ],
    },
    {
        "num":   "H2",
        "title": "Correspondencia entre el Sistema y el Mundo Real",
        "desc":  "El sistema debe usar el idioma del usuario, con palabras, frases y conceptos familiares "
                 "para el usuario en lugar de términos técnicos orientados al sistema.",
        "problemas": [
            {
                "problema":        "Términos técnicos como 'Confianza del modelo', 'Top-3', 'índice de clase' "
                                   "aparecen directamente en la interfaz sin explicación para usuarios agricultores "
                                   "sin formación técnica.",
                "area":            "Resultados del diagnóstico CNN",
                "severidad":       "3 – Mayor",
                "recomendacion":   "Reemplazar o complementar terminología técnica con lenguaje natural. "
                                   "Por ejemplo: 'Nivel de certeza: Alto (93%)' en lugar de 'Confianza: 0.93'. "
                                   "Incluir tooltips explicativos.",
            },
            {
                "problema":        "Los mensajes de error del sistema están en inglés técnico en algunos casos "
                                   "(errores de Flask no personalizados).",
                "area":            "Manejo de errores del sistema",
                "severidad":       "2 – Menor",
                "recomendacion":   "Personalizar todas las páginas de error (400, 403, 404, 500) con mensajes "
                                   "amigables en español y con guía de qué hacer a continuación.",
            },
        ],
    },
    {
        "num":   "H3",
        "title": "Control y Libertad del Usuario",
        "desc":  "Los usuarios eligen funciones del sistema por error y necesitarán una 'salida de emergencia' "
                 "claramente marcada para abandonar el estado no deseado.",
        "problemas": [
            {
                "problema":        "No existe opción para cancelar el proceso de análisis de imagen una vez "
                                   "iniciado. El usuario debe esperar hasta que el modelo CNN termine el procesamiento.",
                "area":            "Motor de Diagnóstico IA",
                "severidad":       "2 – Menor",
                "recomendacion":   "Agregar un botón 'Cancelar análisis' que interrumpa la solicitud HTTP "
                                   "en curso y lleve al usuario de vuelta a la pantalla inicial.",
            },
            {
                "problema":        "Al eliminar un diagnóstico del historial, no se solicita confirmación. "
                                   "La acción es irreversible sin previo aviso.",
                "area":            "Historial de Diagnósticos",
                "severidad":       "3 – Mayor",
                "recomendacion":   "Implementar un diálogo modal de confirmación antes de eliminar registros. "
                                   "Considerar una papelera temporal de 30 días.",
            },
        ],
    },
    {
        "num":   "H4",
        "title": "Consistencia y Estándares",
        "desc":  "Los usuarios no deberían tener que preguntarse si diferentes palabras, situaciones o acciones "
                 "significan lo mismo. Seguir las convenciones de la plataforma.",
        "problemas": [
            {
                "problema":        "El diseño de botones varía en distintas secciones del sistema: algunos usan "
                                   "estilo Bootstrap, otros tienen estilos CSS personalizados inconsistentes "
                                   "(colores, tamaños y tipografía distintos).",
                "area":            "Interfaz general – Componentes UI",
                "severidad":       "2 – Menor",
                "recomendacion":   "Definir y aplicar un sistema de diseño (design system) unificado con "
                                   "componentes reutilizables y paleta de colores consistente en todo el sistema.",
            },
            {
                "problema":        "Las etiquetas de menú de navegación no son uniformes: en algunas páginas "
                                   "aparece 'Diagnóstico' y en otras 'Motor IA', refiriéndose a la misma funcionalidad.",
                "area":            "Menú de navegación",
                "severidad":       "1 – Cosmético",
                "recomendacion":   "Estandarizar todos los términos de navegación a lo largo del sistema. "
                                   "Usar únicamente 'Motor de Diagnóstico' como etiqueta principal.",
            },
        ],
    },
    {
        "num":   "H5",
        "title": "Prevención de Errores",
        "desc":  "Mejor que buenos mensajes de error es un diseño cuidadoso que evite que los problemas ocurran.",
        "problemas": [
            {
                "problema":        "El formulario de carga de imágenes permite seleccionar archivos de cualquier "
                                   "tipo (PDF, DOC, etc.). La validación ocurre solo en el servidor, mostrando "
                                   "el error después de que el usuario ya subió el archivo.",
                "area":            "Formulario de carga de imagen – Motor IA",
                "severidad":       "3 – Mayor",
                "recomendacion":   "Implementar validación del tipo y tamaño de archivo en el lado del cliente "
                                   "(JavaScript) antes del envío. Mostrar los formatos aceptados (JPG, PNG, WEBP) "
                                   "claramente junto al campo de carga.",
            },
            {
                "problema":        "No existe límite visible de tamaño de archivo en la interfaz. El usuario "
                                   "puede intentar subir imágenes muy grandes causando timeouts sin explicación.",
                "area":            "Carga de imágenes",
                "severidad":       "2 – Menor",
                "recomendacion":   "Indicar claramente el límite de tamaño junto al botón de carga. "
                                   "Implementar validación de tamaño en cliente con mensaje preventivo.",
            },
        ],
    },
    {
        "num":   "H6",
        "title": "Reconocimiento antes que Recuerdo",
        "desc":  "Minimizar la carga de memoria del usuario haciendo que los objetos, acciones y opciones sean visibles.",
        "problemas": [
            {
                "problema":        "En la pantalla de resultados del diagnóstico, los nombres científicos de las "
                                   "enfermedades no siempre incluyen el nombre común en español. El usuario debe "
                                   "recordar o buscar externamente qué significa cada término.",
                "area":            "Pantalla de resultados del diagnóstico",
                "severidad":       "2 – Menor",
                "recomendacion":   "Mostrar siempre el nombre común en español prominentemente y el nombre "
                                   "científico como información secundaria. Incluir ícono de información (ℹ) "
                                   "con descripción expandible.",
            },
            {
                "problema":        "No existe una guía de inicio rápido o tutorial contextual para usuarios "
                                   "nuevos que explique cómo usar el motor de diagnóstico.",
                "area":            "Onboarding de nuevos usuarios",
                "severidad":       "2 – Menor",
                "recomendacion":   "Implementar un tour guiado de primer uso (usando bibliotecas como intro.js) "
                                   "o un panel de ayuda contextual siempre visible.",
            },
        ],
    },
    {
        "num":   "H7",
        "title": "Flexibilidad y Eficiencia de Uso",
        "desc":  "Los aceleradores —invisibles para el usuario novato— pueden acelerar la interacción del usuario "
                 "experto, de modo que el sistema pueda atender tanto a usuarios sin experiencia como con experiencia.",
        "problemas": [
            {
                "problema":        "No existen atajos de teclado ni funciones avanzadas para usuarios frecuentes "
                                   "como técnicos agrícolas que usan el sistema diariamente.",
                "area":            "Funcionalidades generales",
                "severidad":       "1 – Cosmético",
                "recomendacion":   "Implementar atajos de teclado para acciones frecuentes (Ej: Ctrl+U para "
                                   "cargar imagen, Ctrl+H para historial). Agregar modo avanzado con filtros "
                                   "adicionales para usuarios expertos.",
            },
            {
                "problema":        "El historial de diagnósticos no permite exportar ni filtrar por fecha, tipo "
                                   "de enfermedad o cultivo. Los usuarios expertos necesitan estas capacidades "
                                   "para gestionar múltiples diagnósticos.",
                "area":            "Historial de Diagnósticos",
                "severidad":       "2 – Menor",
                "recomendacion":   "Implementar filtros por fecha, tipo de enfermedad y cultivo. Agregar "
                                   "funcionalidad de exportación en CSV/PDF para técnicos agrícolas.",
            },
        ],
    },
    {
        "num":   "H8",
        "title": "Diseño Estético y Minimalista",
        "desc":  "Los diálogos no deben contener información irrelevante o raramente necesaria. Cada unidad "
                 "de información adicional en un diálogo compite con la información relevante.",
        "problemas": [
            {
                "problema":        "La pantalla de resultados muestra simultáneamente: imagen diagnósticada, "
                                   "3 diagnósticos con porcentajes, descripción de enfermedad, tratamiento, "
                                   "suplemento, chat IA y calificación. La sobrecarga visual puede abrumar "
                                   "al agricultor.",
                "area":            "Pantalla de resultados del diagnóstico",
                "severidad":       "2 – Menor",
                "recomendacion":   "Reorganizar la información usando un diseño progresivo: mostrar primero "
                                   "el diagnóstico principal y la acción recomendada; detalles adicionales "
                                   "en pestañas o secciones expandibles.",
            },
            {
                "problema":        "Algunas páginas muestran mensajes flash informativos que no son relevantes "
                                   "para la tarea actual del usuario.",
                "area":            "Mensajes flash del sistema",
                "severidad":       "1 – Cosmético",
                "recomendacion":   "Mostrar solo mensajes flash pertinentes a la acción inmediatamente realizada. "
                                   "Implementar un sistema de notificaciones categorizado (éxito, error, info).",
            },
        ],
    },
    {
        "num":   "H9",
        "title": "Ayuda a los Usuarios a Reconocer, Diagnosticar y Recuperarse de los Errores",
        "desc":  "Los mensajes de error deben expresarse en lenguaje sencillo (sin códigos), indicar con precisión "
                 "el problema y sugerir constructivamente una solución.",
        "problemas": [
            {
                "problema":        "Cuando el modelo CNN detecta una imagen borrosa y la rechaza, el mensaje "
                                   "de error solo dice 'Imagen borrosa detectada' sin guiar al usuario sobre "
                                   "cómo tomar una mejor foto.",
                "area":            "Validación de calidad de imagen",
                "severidad":       "3 – Mayor",
                "recomendacion":   "Complementar el mensaje de error con consejos prácticos: 'Asegúrese de "
                                   "que la hoja esté bien iluminada y a 20-30 cm de la cámara' y mostrar "
                                   "un ejemplo visual de imagen válida.",
            },
            {
                "problema":        "Los errores de autenticación (contraseña incorrecta, email no registrado) "
                                   "usan mensajes genéricos que no distinguen el tipo de error por razones "
                                   "de seguridad, pero no ofrecen alternativas de recuperación.",
                "area":            "Módulo de autenticación",
                "severidad":       "2 – Menor",
                "recomendacion":   "Agregar un enlace prominente de '¿Olvidaste tu contraseña?' en el formulario "
                                   "de login y un enlace a la página de registro cuando el email no está registrado.",
            },
        ],
    },
    {
        "num":   "H10",
        "title": "Ayuda y Documentación",
        "desc":  "Aunque es mejor si el sistema puede usarse sin documentación, puede ser necesario proporcionar "
                 "ayuda y documentación fácil de buscar, enfocada en la tarea del usuario.",
        "problemas": [
            {
                "problema":        "No existe una sección de preguntas frecuentes (FAQ) ni documentación de "
                                   "ayuda integrada en el sistema. Los usuarios deben consultar documentación "
                                   "externa para entender el funcionamiento del motor IA.",
                "area":            "Sistema general – Documentación integrada",
                "severidad":       "2 – Menor",
                "recomendacion":   "Implementar una sección 'Ayuda' con FAQ sobre: cómo tomar buenas fotos, "
                                   "cómo interpretar los resultados, qué hacer ante cada diagnóstico. "
                                   "Integrar búsqueda dentro de la ayuda.",
            },
            {
                "problema":        "El asistente AndesGPT no tiene instrucciones claras sobre sus capacidades "
                                   "y limitaciones. Los usuarios pueden no saber qué preguntas puede responder.",
                "area":            "Asistente IA AndesGPT",
                "severidad":       "1 – Cosmético",
                "recomendacion":   "Mostrar ejemplos de preguntas sugeridas al abrir el chat por primera vez. "
                                   "Incluir una breve descripción de las capacidades del asistente.",
            },
        ],
    },
]

# ─────────────────────────────────────────
# RENDER EACH HEURÍSTICA
# ─────────────────────────────────────────
for h in heuristicas:
    # Heurística title
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(f"  {h['num']} – {h['title']}")
    r.font.name  = "Times New Roman"
    r.font.size  = Pt(13)
    r.font.bold  = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    # Add shading to the paragraph via XML
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  '004680')
    pPr.append(shd)

    # Description
    add_paragraph(doc,
        f"Principio: {h['desc']}",
        font_size=11, italic=True, color=(60,60,60),
        space_before=4, space_after=8)

    # Problems table
    for idx, prob in enumerate(h['problemas'], 1):
        # Problem header
        p_hdr = doc.add_paragraph()
        p_hdr.paragraph_format.space_before = Pt(6)
        p_hdr.paragraph_format.space_after  = Pt(2)
        r_hdr = p_hdr.add_run(f"  Problema {idx}")
        r_hdr.font.name  = "Times New Roman"
        r_hdr.font.size  = Pt(11)
        r_hdr.font.bold  = True
        r_hdr.font.color.rgb = RGBColor(255, 255, 255)
        pPr2 = p_hdr._p.get_or_add_pPr()
        shd2 = OxmlElement('w:shd')
        shd2.set(qn('w:val'),   'clear')
        shd2.set(qn('w:color'), 'auto')
        shd2.set(qn('w:fill'),  '0064A0')
        pPr2.append(shd2)

        # Details table
        tbl2 = doc.add_table(rows=4, cols=2)
        tbl2.style = 'Table Grid'

        labels = ["Problema Identificado", "Área del Sistema", "Severidad", "Recomendación"]
        values = [prob['problema'], prob['area'], prob['severidad'], prob['recomendacion']]
        fill_label = "D6E4F0"
        fill_value = "FFFFFF"
        fill_alt   = "EEF5FB"

        for row_i, (lbl, val) in enumerate(zip(labels, values)):
            row = tbl2.rows[row_i]
            row.cells[0].width = Cm(4.5)
            row.cells[1].width = Cm(13)
            shade_cell(row.cells[0], fill_label)
            shade_cell(row.cells[1], fill_alt if row_i % 2 == 0 else fill_value)
            cell_text(row.cells[0], lbl, bold=True, size=10, color=(0,70,127))
            # severity color
            if lbl == "Severidad":
                sev_num = int(val[0]) if val[0].isdigit() else 0
                sev_colors = {0:(80,80,80), 1:(150,120,0), 2:(200,100,0), 3:(180,30,0), 4:(120,0,0)}
                cell_text(row.cells[1], val, bold=True, size=10,
                          color=sev_colors.get(sev_num, (30,30,30)))
            else:
                cell_text(row.cells[1], val, size=10)

        add_paragraph(doc, "", space_after=4)

add_page_break(doc)

# ═══════════════════════════════════════════════════
#  RESUMEN CUANTITATIVO
# ═══════════════════════════════════════════════════
add_heading_custom(doc, "VII. RESUMEN CUANTITATIVO DE PROBLEMAS", 1)

add_paragraph(doc,
    "La siguiente tabla consolida el número y la severidad de los problemas encontrados por heurística:",
    font_size=12, space_after=8)

summary_data = [
    ("Heurística", "N° Problemas", "Severidad Máxima", "Prioridad"),
    ("H1 – Visibilidad del Estado",           "2", "3 – Mayor",      "Alta"),
    ("H2 – Correspondencia Mundo Real",       "2", "3 – Mayor",      "Alta"),
    ("H3 – Control y Libertad",               "2", "3 – Mayor",      "Alta"),
    ("H4 – Consistencia y Estándares",        "2", "2 – Menor",      "Media"),
    ("H5 – Prevención de Errores",            "2", "3 – Mayor",      "Alta"),
    ("H6 – Reconocimiento antes que Recuerdo","2", "2 – Menor",      "Media"),
    ("H7 – Flexibilidad y Eficiencia",        "2", "2 – Menor",      "Media"),
    ("H8 – Diseño Estético y Minimalista",    "2", "2 – Menor",      "Media"),
    ("H9 – Recuperación de Errores",          "2", "3 – Mayor",      "Alta"),
    ("H10 – Ayuda y Documentación",           "2", "2 – Menor",      "Media"),
    ("TOTAL",                                 "20","—",              "—"),
]

tbl3 = doc.add_table(rows=len(summary_data), cols=4)
tbl3.style = 'Table Grid'
col_w = [Cm(7.5), Cm(2.8), Cm(3.8), Cm(3)]

for i, row_data in enumerate(summary_data):
    row = tbl3.rows[i]
    is_hdr   = (i == 0)
    is_total = (i == len(summary_data) - 1)
    for j, (val, w) in enumerate(zip(row_data, col_w)):
        cell = row.cells[j]
        cell.width = w
        if is_hdr:
            shade_cell(cell, "004680")
            cell_text(cell, val, bold=True, size=10, color=(255,255,255),
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        elif is_total:
            shade_cell(cell, "D6E4F0")
            cell_text(cell, val, bold=True, size=10, color=(0,70,127),
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            shade_cell(cell, "EEF5FB" if i % 2 == 0 else "FFFFFF")
            # Color severity
            if j == 2:
                sev_map = {"3 – Mayor":(180,30,0), "2 – Menor":(200,100,0), "1 – Cosmético":(150,120,0)}
                clr = sev_map.get(val, (30,30,30))
                cell_text(cell, val, bold=True, size=10, color=clr,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
            elif j == 3:
                pri_map = {"Alta":(180,30,0), "Media":(200,100,0)}
                clr = pri_map.get(val, (30,30,30))
                cell_text(cell, val, bold=True, size=10, color=clr,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                cell_text(cell, val, size=10,
                          align=WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT)

add_paragraph(doc, "", space_after=8)

# ═══════════════════════════════════════════════════
#  CONCLUSIONES
# ═══════════════════════════════════════════════════
add_heading_custom(doc, "VIII. CONCLUSIONES", 1)

conclusiones = [
    ("1.", "Se identificaron un total de 20 problemas de usabilidad distribuidos en las 10 heurísticas de "
           "Nielsen, de los cuales 10 corresponden a severidad Mayor (nivel 3), requiriendo atención prioritaria "
           "antes del despliegue en producción."),
    ("2.", "Las heurísticas con mayor impacto negativo en la experiencia del usuario son: H1 (Visibilidad del "
           "Estado), H3 (Control y Libertad), H5 (Prevención de Errores) y H9 (Recuperación de Errores), todas "
           "con problemas de severidad Mayor que afectan directamente la confianza del agricultor en el sistema."),
    ("3.", "El diseño actual de PlantAndes es funcionally robusto desde el punto de vista técnico, pero requiere "
           "mejoras significativas en la comunicación con usuarios no técnicos, especialmente considerando que "
           "su público objetivo son agricultores de zonas andinas con posible baja experiencia digital."),
    ("4.", "Se recomienda priorizar la implementación de indicadores de progreso, mensajes de error claros con "
           "guía de acción, y una validación de imágenes en el lado del cliente como mejoras de alto impacto "
           "y bajo costo de implementación."),
    ("5.", "El soporte multilenguaje Español/Quechua es un punto fuerte del sistema y debería extenderse también "
           "a los mensajes de error y la documentación de ayuda integrada para maximizar la accesibilidad "
           "en las comunidades objetivo de Cusco."),
]

for num, text in conclusiones:
    p = doc.add_paragraph()
    p.paragraph_format.space_before  = Pt(4)
    p.paragraph_format.space_after   = Pt(6)
    p.paragraph_format.left_indent   = Cm(0.5)
    r1 = p.add_run(f"{num} ")
    set_font(r1, bold=True, size=12)
    r2 = p.add_run(text)
    set_font(r2, size=12)

add_paragraph(doc, "", space_after=10)

# ═══════════════════════════════════════════════════
#  RECOMENDACIONES GENERALES
# ═══════════════════════════════════════════════════
add_heading_custom(doc, "IX. RECOMENDACIONES GENERALES", 1)

recomendaciones = [
    ("Corto plazo (Sprint inmediato):",
     "Implementar indicadores de carga durante el procesamiento de imagen, añadir diálogo de confirmación "
     "antes de eliminar diagnósticos, y agregar validación de tipo/tamaño de archivo en el cliente."),
    ("Mediano plazo (1-2 Sprints):",
     "Desarrollar un sistema de diseño unificado, crear una sección de ayuda con FAQ integrado, y mejorar "
     "los mensajes de error del sistema de validación de imágenes con guías visuales."),
    ("Largo plazo (Versión 2.0):",
     "Implementar un tour de onboarding para nuevos usuarios, agregar funcionalidades de exportación y "
     "filtrado avanzado en el historial, y desarrollar atajos de teclado para usuarios frecuentes."),
]

for titulo, texto in recomendaciones:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.5)
    r1 = p.add_run(f"{titulo} ")
    set_font(r1, bold=True, size=12, color=(0,70,127))
    r2 = p.add_run(texto)
    set_font(r2, size=12)

add_paragraph(doc, "", space_after=10)

# ═══════════════════════════════════════════════════
#  REFERENCIAS
# ═══════════════════════════════════════════════════
add_heading_custom(doc, "X. REFERENCIAS", 1)

refs = [
    "Nielsen, J. (1994). Heuristic Evaluation. In Nielsen, J., and Mack, R.L. (Eds.), "
    "Usability Inspection Methods. John Wiley & Sons.",
    "Nielsen, J., & Molich, R. (1990). Heuristic evaluation of user interfaces. "
    "Proceedings of the ACM CHI'90 Conference, 249-256.",
    "Nielsen, J. (1995). How to Rate the Severity of Usability Problems. "
    "Nielsen Norman Group. https://www.nngroup.com/articles/how-to-rate-the-severity-of-usability-problems/",
    "Norman, D. A. (2013). The Design of Everyday Things: Revised and Expanded Edition. Basic Books.",
    "ISO 9241-11 (2018). Ergonomics of Human-System Interaction – Usability: Definitions and Concepts. "
    "International Organization for Standardization.",
]

for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before    = Pt(4)
    p.paragraph_format.space_after     = Pt(4)
    p.paragraph_format.left_indent     = Cm(1)
    p.paragraph_format.first_line_indent = Cm(-1)
    r = p.add_run(f"[{i}] {ref}")
    set_font(r, size=11)

# ─────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────
out_path = r'c:\Users\PC\Desktop\IDEA-PROYECTO\Evaluacion_Heuristica_Nielsen_PlantAndes.docx'
doc.save(out_path)
print(f"Documento guardado exitosamente en:\n{out_path}")
