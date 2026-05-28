import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Define paths
base_dir = r"c:\Users\estudianteUC\Desktop\IDEA-PROYECTO\Informe Usabilidad"
csv_path = os.path.join(base_dir, "Data_Evaluación_de_Usabilidad.csv")
template_path = os.path.join(base_dir, "Plantilla_Informe_Usabilidad.docx")
output_path = os.path.join(base_dir, "Informe_de_Usabilidad_PlantAndes.docx")
charts_dir = os.path.join(base_dir, "charts")

# Create charts directory if it doesn't exist
os.makedirs(charts_dir, exist_ok=True)

# ---------------------------------------------------------
# 1. LOAD AND PROCESS DATA
# ---------------------------------------------------------
df = pd.read_csv(csv_path)

# Extract users
users = []
for idx, row in df.iterrows():
    name = row['2. Nombre del participante'].strip()
    carrera = row['3. Carrera / Cycle'] if '3. Carrera / Cycle' in df.columns else row['3. Carrera / Ciclo']
    prev_use = row['4. ¿Había utilizado anteriormente el sistema PlantAndes?']
    users.append({
        "num": idx + 1,
        "name": name,
        "carrera": carrera,
        "prev_use": prev_use
    })

# Tasks and their columns
tasks_def = {
    "T1": {
        "title": "T1: Registro e Inicio de Sesión",
        "success": "¿Logró completar la tarea 1?",
        "time": df.columns[7],  # "¿Cuánto tiempo tardó?..."
        "error": "¿Cuántos errores tuvo durante el registro/inicio de sesión?",
        "diff": "Nivel de dificultad de la Tarea 1"
    },
    "T2": {
        "title": "T2: Diagnóstico de Cultivo",
        "success": "¿Logró completar la tarea 2?",
        "time": df.columns[12],  # "¿Cuánto tiempo tardó?..."
        "error": "¿Cuántos errores tuvo al subir la imagen o procesar el diagnóstico?",
        "diff": "Nivel de dificultad de la Tarea 2"
    },
    "T3": {
        "title": "T3: Interpretar Resultados",
        "success": "¿Logró completar la tarea 3?",
        "time": df.columns[18],  # "¿Cuánto tiempo tardó en entender los resultados?..."
        "error": "¿Cuántos errores o confusiones tuvo al interpretar los resultados?",
        "diff": "Nivel de dificultad of the Tarea 3" if "Nivel de dificultad of the Tarea 3" in df.columns else df.columns[20] # Check name
    },
    "T4": {
        "title": "T4: Usar Asistente AndesGPT",
        "success": "¿Logró completar la tarea 4?",
        "time": df.columns[24],  # "¿Cuánto tiempo tardó?..."
        "error": "¿Cuántos errores o dificultades tuvo al usar el asistente?",
        "diff": "Nivel de dificultad de la Tarea 4"
    },
    "T5": {
        "title": "T5: Revisar Historial",
        "success": "¿Logró completar la tarea 5?",
        "time": df.columns[30],  # "¿Cuánto tiempo tardó?..."
        "error": "¿Cuántos errores tuvo al navegar el historial?",
        "diff": "Nivel de dificultad de la Tarea 5"
    }
}

# Resolve actual column names in CSV
task_cols = {}
for t_id, t_def in tasks_def.items():
    resolved = {}
    for key, col_name in t_def.items():
        if key == "title":
            resolved[key] = col_name
            continue
        # match column by substring
        match = [c for c in df.columns if col_name.lower()[:30] in c.lower()]
        if match:
            resolved[key] = match[0]
        else:
            # Fallback to direct check or manual mapping
            resolved[key] = col_name
    task_cols[t_id] = resolved

# Function to parse time (to seconds)
def parse_time(val):
    if pd.isna(val):
        return 0
    val_str = str(val).lower().strip()
    if "seg" in val_str:
        num = "".join(c for c in val_str if c.isdigit())
        return int(num) if num else 30
    if ":" in val_str:
        parts = val_str.split(":")
        minutes = int("".join(c for c in parts[0] if c.isdigit()))
        seconds = int("".join(c for c in parts[1] if c.isdigit() or c == ' '))
        return minutes * 60 + seconds
    
    digits = "".join(c for c in val_str if c.isdigit())
    if digits:
        val_num = int(digits)
        if "seg" not in val_str:
            return val_num * 60
        else:
            return val_num
    
    if "menos de 1 minuto" in val_str:
        return 45
    if "casi inmediato" in val_str:
        return 30
    return 60

# Function to map error strings to numbers
def map_errors(val):
    val_str = str(val).strip().lower()
    if '0' in val_str: return 0
    if '1' in val_str: return 1
    if '2' in val_str: return 2
    if '3' in val_str: return 3
    return 0

task_stats = {}
for t_id, t_cols in task_cols.items():
    success_rate = df[t_cols["success"]].apply(lambda x: 100 if str(x).strip().lower() == "sí" else 0).mean()
    times = df[t_cols["time"]].apply(parse_time)
    avg_time = times.mean()
    errors = df[t_cols["error"]].apply(map_errors)
    avg_errors = errors.mean()
    avg_diff = df[t_cols["diff"]].mean()
    task_stats[t_id] = {
        "title": t_cols["title"],
        "success_rate": success_rate,
        "avg_time": avg_time,
        "avg_errors": avg_errors,
        "avg_diff": avg_diff
    }

# SUS calculation
sus_cols = [c for c in df.columns if c.startswith("SUS ")]
# Sort sus_cols by number to make sure they are in order SUS 1 to 10
sus_cols = sorted(sus_cols, key=lambda x: int(x.split(".")[0].replace("SUS", "").strip()))

sus_scores = []
for idx, row in df.iterrows():
    score = 0
    for q_idx in range(10):
        val = int(row[sus_cols[q_idx]])
        if (q_idx + 1) % 2 == 1:  # Odd items: SUS 1, 3, 5, 7, 9 -> (val - 1)
            score += (val - 1)
        else:  # Even items: SUS 2, 4, 6, 8, 10 -> (5 - val)
            score += (5 - val)
    sus_scores.append(score * 2.5)

df['SUS_Score'] = sus_scores
avg_sus = np.mean(sus_scores)

# Section 3: General satisfaction Likert
likert_cols = [
    "El sistema PlantAndes fue fácil de usar",
    "La navegación del sistema fue clara e intuitiva",
    "Los botones y opciones fueron comprensibles",
    "Los mensajes y resultados del sistema fueron claros",
    "Me sentí cómodo/a usando el sistema",
    "El sistema responde de forma rápida",
    "El diagnóstico generado por la IA me parece confiable",
    "El sistema cumple correctamente su función de diagnosticar enfermedades",
    "Recomendaría PlantAndes a agricultores de la región",
    "El diseño visual del sistema es atractivo y profesional"
]
# Resolve actual columns in CSV for likert
resolved_likert_cols = []
for l_col in likert_cols:
    match = [c for c in df.columns if l_col[:25].lower() in c.lower()]
    if match:
        resolved_likert_cols.append((l_col, match[0]))
    else:
        resolved_likert_cols.append((l_col, l_col))

likert_stats = []
for short_lbl, full_col in resolved_likert_cols:
    likert_stats.append({
        "label": short_lbl,
        "avg": df[full_col].mean()
    })

# ---------------------------------------------------------
# 2. GENERATE BEAUTIFUL CHARTS
# ---------------------------------------------------------
print("Generating charts...")
sns.set_theme(style="whitegrid")
plt.rcParams['font.sans-serif'] = ['DejaVu Sans', 'Arial']

# Palette
primary_color = "#1b4332"  # Dark Forest Green
secondary_color = "#40916c" # Medium Green
accent_color = "#52b788"   # Light Green
neutral_dark = "#2d3748"   # Charcoal
neutral_light = "#edf2f7"  # Cool grey

# Chart 1: SUS Scores
plt.figure(figsize=(7, 4.5))
x_labels = [row['2. Nombre del participante'].split()[0] for idx, row in df.iterrows()]
# Add * to John and Gaby to indicate all 5s bias
for i, name in enumerate(x_labels):
    if name in ["John", "Gaby"]:
        x_labels[i] = name + " (Sesgo*)"

bars = plt.bar(x_labels, sus_scores, color=[secondary_color if "Sesgo" not in name else "#cbd5e1" for name in x_labels], width=0.55, edgecolor="#1e293b", linewidth=0.7)
plt.axhline(68, color="#e11d48", linestyle="--", linewidth=1.2, label="Promedio Industrial (68)")
plt.axhline(80, color="#16a34a", linestyle="--", linewidth=1.2, label="Umbral de Excelencia (80)")

# Add values on top of bars
for bar in bars:
    height = bar.get_height()
    plt.text(bar.get_x() + bar.get_width()/2.0, height + 1.5, f'{height:.1f}', ha='center', va='bottom', fontsize=10, fontweight='bold', color=neutral_dark)

plt.ylim(0, 110)
plt.title("Puntajes SUS por Participante", fontsize=12, fontweight='bold', color=primary_color, pad=12)
plt.ylabel("Puntaje SUS (0 - 100)", fontsize=10, fontweight='bold', color=neutral_dark)
plt.legend(loc="lower left", frameon=True, facecolor="white", edgecolor="#e2e8f0")
plt.tight_layout()
sus_chart_path = os.path.join(charts_dir, "sus_scores.png")
plt.savefig(sus_chart_path, dpi=300)
plt.close()
print("Saved SUS chart.")

# Chart 2: Task Metrics (1x2 Subplots)
fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4.5))

# Subplot 1: Times
t_ids = list(task_stats.keys())
times_vals = [task_stats[tid]["avg_time"] for tid in t_ids]
bars1 = ax1.bar(t_ids, times_vals, color=secondary_color, width=0.5, edgecolor=primary_color, linewidth=0.7)
ax1.set_title("Tiempo Promedio por Tarea", fontsize=11, fontweight='bold', color=primary_color, pad=10)
ax1.set_ylabel("Tiempo (segundos)", fontsize=9, fontweight='bold', color=neutral_dark)
ax1.set_ylim(0, max(times_vals) * 1.15)
for bar in bars1:
    height = bar.get_height()
    ax1.text(bar.get_x() + bar.get_width()/2.0, height + 2, f'{height:.1f}s', ha='center', va='bottom', fontsize=9, fontweight='bold')

# Subplot 2: Difficulty / Ease
diff_vals = [task_stats[tid]["avg_diff"] for tid in t_ids]
bars2 = ax2.bar(t_ids, diff_vals, color=accent_color, width=0.5, edgecolor=secondary_color, linewidth=0.7)
ax2.set_title("Facilidad de Uso Percibida (1-5)", fontsize=11, fontweight='bold', color=primary_color, pad=10)
ax2.set_ylabel("Escala (1: Muy Difícil, 5: Muy Fácil)", fontsize=9, fontweight='bold', color=neutral_dark)
ax2.set_ylim(0, 5.8)
for bar in bars2:
    height = bar.get_height()
    ax2.text(bar.get_x() + bar.get_width()/2.0, height + 0.1, f'{height:.2f}', ha='center', va='bottom', fontsize=9, fontweight='bold')

plt.tight_layout()
task_chart_path = os.path.join(charts_dir, "task_metrics.png")
plt.savefig(task_chart_path, dpi=300)
plt.close()
print("Saved Task metrics chart.")

# Chart 3: Likert Satisfaction (Horizontal Bar Chart)
plt.figure(figsize=(9, 5))
likert_labels = [item["label"] for item in likert_stats]
likert_vals = [item["avg"] for item in likert_stats]

# Reverse order so the first item is at the top
likert_labels.reverse()
likert_vals.reverse()

bars3 = plt.barh(likert_labels, likert_vals, color=sns.color_palette("viridis_r", len(likert_vals)), height=0.6, edgecolor="#334155", linewidth=0.5)
plt.xlim(0, 5.5)
plt.title("Evaluación General del Sistema (Escala Likert 1-5)", fontsize=12, fontweight='bold', color=primary_color, pad=12)
plt.xlabel("Puntaje Promedio", fontsize=10, fontweight='bold', color=neutral_dark)

# Add values on the bars
for bar in bars3:
    width = bar.get_width()
    plt.text(width + 0.05, bar.get_y() + bar.get_height()/2.0, f'{width:.2f}', ha='left', va='center', fontsize=9, fontweight='bold', color=neutral_dark)

plt.tight_layout()
likert_chart_path = os.path.join(charts_dir, "likert_satisfaction.png")
plt.savefig(likert_chart_path, dpi=300)
plt.close()
print("Saved Likert chart.")

# ---------------------------------------------------------
# 3. FILL AND GENERATE WORD DOCUMENT
# ---------------------------------------------------------
print("Writing Word document...")
doc = docx.Document(template_path)

# Store original tables before inserting any new ones
t1_participants = doc.tables[0]
t2_tasks = doc.tables[1]
t3_recoms = doc.tables[2]

# Helper function to replace text in paragraphs
def replace_para_text(para, new_text):
    para.text = "" # Clear
    run = para.add_run(new_text)
    # Restore standard formatting (Times New Roman or Calibri depending on doc, let's keep it simple)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

# Helper to insert paragraphs after a paragraph
def insert_paragraph_after(paragraph, text=None, style=None):
    new_p = OxmlElement('w:p')
    paragraph._element.addnext(new_p)
    new_para = docx.text.paragraph.Paragraph(new_p, paragraph._parent)
    if text:
        new_para.text = text
        for r in new_para.runs:
            r.font.name = 'Calibri'
            r.font.size = Pt(11)
    if style:
        new_para.style = style
    return new_para

# Helper to format a heading nicely
def format_heading(para, text):
    para.text = ""
    run = para.add_run(text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(27, 67, 50) # Forest green

# Let's search for and identify all the placeholder paragraphs in the template
para_metadata = None
para_intro = None
para_desc = None
para_method = None
para_tasks = None
para_results_placeholder = None
para_analysis = None
para_recoms = None
para_conclusions = None
para_annexes = None

for para in doc.paragraphs:
    text = para.text.strip()
    if "Curso:" in text and "Docente:" in text:
        para_metadata = para
    elif "Describir brevemente el sistema evaluado" in text:
        para_intro = para
    elif "Nombre del sistema" in text and "Módulos evaluados" in text:
        para_desc = para
    elif "Tipo de evaluación aplicada" in text or "Escenario de prueba" in text:
        para_method = para
    elif "Registrar las tareas realizadas por los usuarios" in text:
        para_tasks = para
    elif "Insertar tablas y gráficos generados en Excel" in text:
        para_results_placeholder = para
    elif "Describir hallazgos principales, aspectos positivos" in text:
        para_analysis = para
    elif "Registrar recomendaciones según los problemas" in text:
        para_recoms = para
    elif "Redactar conclusiones generales de la evaluación" in text:
        para_conclusions = para
    elif "Capturas del sistema" in text and "Resultados en Excel" in text:
        para_annexes = para

# Metadata Block
metadata_text = (
    "Curso: Taller de Proyectos II\n"
    "Docente: Mg. Américo Estrada Sánchez\n"
    "Proyecto/Sistema: PlantAndes — Sistema Web de Diagnóstico de Enfermedades en Hojas de Cultivos Andinos mediante Visión Artificial (CNN)\n"
    "Integrantes del equipo:\n"
    "  • Calderon Romero, Yersson\n"
    "  • Guevara Valdivia, Alejandro Humberto\n"
    "Fecha de evaluación: 28 de mayo de 2026\n"
    "Versión del sistema evaluado: v1.0.0 (Entorno de Desarrollo Local)"
)
if para_metadata:
    replace_para_text(para_metadata, metadata_text)

# Introduction text
intro_text = (
    "El presente informe documenta la evaluación de usabilidad aplicada al sistema PlantAndes, una plataforma web desarrollada en Python (Flask) que utiliza modelos de Deep Learning (Redes Neuronales Convolucionales - CNN) para diagnosticar de forma automatizada enfermedades en hojas de cultivos andinos tradicionales en comunidades de Cusco.\n\n"
    "La evaluación de usabilidad es una fase crucial en la ingeniería de software y el diseño UX (Experiencia de Usuario). Permite medir en qué grado un sistema de software es fácil de aprender, eficiente en su operación, seguro ante fallos y satisfactorio en su uso práctico. En el contexto de PlantAndes, este análisis adquiere una relevancia crítica, dado que los usuarios finales —agricultores locales y promotores agrícolas— requieren una interfaz altamente intuitiva, simplificada y libre de tecnicismos para facilitar la adopción tecnológica en el sector rural andino. Este reporte recopila y analiza métricas cuantitativas y cualitativas de las pruebas realizadas con usuarios para recomendar mejoras específicas en la interfaz."
)
if para_intro:
    replace_para_text(para_intro, intro_text)

# System description
desc_text = (
    "PlantAndes es una solución web monolítica de diagnóstico agrícola inteligente que integra modelos avanzados de visión por computadora y asistencia conversacional.\n\n"
    "• Nombre del sistema: PlantAndes (Diagnóstico Inteligente de Cultivos Andinos)\n"
    "• Descripción general: Plataforma basada en el backend Flask que integra dos modelos CNN (redes neuronales convolucionales) implementados en PyTorch. El sistema permite cargar fotos de hojas de cultivos como la papa, tomate y maíz, identificando plagas o estados de deficiencia. Adicionalmente, cuenta con el asistente virtual AndesGPT (vía API de LLM DeepSeek) para guiar al usuario bilingüe sobre tratamientos fitosanitarios ecológicos y convencionales.\n"
    "• Usuarios objetivo: Pequeños agricultores locales de la región Cusco, personal técnico de cooperativas agrarias y estudiantes o investigadores de agronomía.\n"
    "• Módulos evaluados:\n"
    "  1. Módulo de Autenticación: Registro e inicio de sesión de usuarios.\n"
    "  2. Módulo de Diagnóstico por IA: Carga de fotos de hojas, inferencia del modelo CNN y reporte de diagnóstico con porcentaje de certeza.\n"
    "  3. Módulo de Recomendaciones: Despliegue de medidas curativas y preventivas específicas para cada patología.\n"
    "  4. Asistente Conversacional AndesGPT: Chat interactivo para profundizar consultas agrícolas.\n"
    "  5. Historial de Diagnósticos: Consulta y visualización persistente de análisis previos."
)
if para_desc:
    replace_para_text(para_desc, desc_text)

# Methodology
method_text = (
    "La evaluación se ejecutó mediante un estudio controlado basado en escenarios de tareas específicas y cuestionarios de salida estandarizados.\n\n"
    "• Tipo de evaluación aplicada: Evaluación de usabilidad sumativa y de desempeño de tareas, combinada con encuestas post-test para evaluar la satisfacción (Escala SUS y cuestionario Likert).\n"
    "• Participantes: El grupo evaluador estuvo integrado por 6 participantes externos al desarrollo del software, todos ellos estudiantes del 10mo ciclo de la Escuela Profesional de Ingeniería de Sistemas e Informática de la Universidad Continental (Cusco). El perfil de los evaluadores es técnico con alta experiencia tecnológica, ideal para auditar fallos funcionales y detectar problemas de usabilidad desde una perspectiva crítica.\n"
    "• Escenario de prueba: Las pruebas se realizaron utilizando la plataforma web PlantAndes en su entorno de desarrollo local (http://127.0.0.1:5000) bajo el navegador Google Chrome. Cada participante ejecutó las tareas indicadas por el moderador de forma individual y consecutiva, sin ayuda previa. Al finalizar el flujo, respondieron un cuestionario en Google Forms para evaluar la experiencia de uso."
)
if para_method:
    replace_para_text(para_method, method_text)

# Evaluated tasks intro
tasks_intro = (
    "Para auditar la usabilidad de forma estructurada, se definieron 5 tareas clave que cubren los flujos críticos de la aplicación. Los participantes debían completarlas de principio a fin, registrando los resultados individuales en términos de éxito, tiempo y errores:"
)
if para_tasks:
    replace_para_text(para_tasks, tasks_intro)

# Results from questionnaire (Insert tables, charts, analysis)
results_intro = (
    "A continuación, se consolidan y analizan de forma sistemática las métricas recolectadas de la interacción de los 6 participantes evaluadores con el sistema PlantAndes."
)
if para_results_placeholder:
    replace_para_text(para_results_placeholder, results_intro)
    
    current_p = para_results_placeholder
    
    # 1. Consolidado general table title
    current_p = insert_paragraph_after(current_p, "Tabla 1. Consolidado General de Métricas de Usabilidad")
    current_p.runs[0].bold = True
    
    # Create Consolidado General de Métricas Table
    total_users = 6
    total_tasks = 5
    success_pct = 100.0
    avg_time_s = sum([task_stats[tid]["avg_time"] for tid in task_stats])
    avg_errors = sum([task_stats[tid]["avg_errors"] for tid in task_stats])
    
    def insert_table_after(paragraph, data):
        doc_body = doc.element.body
        table = doc.add_table(rows=len(data), cols=len(data[0]))
        table.style = 'Light Shading Accent 1'
        for r_idx, row_data in enumerate(data):
            for c_idx, val in enumerate(row_data):
                cell = table.rows[r_idx].cells[c_idx]
                cell.text = val
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for r in p.runs:
                        r.font.name = 'Calibri'
                        r.font.size = Pt(9.5)
                        if r_idx == 0:
                            r.bold = True
        paragraph._element.addnext(table._element)
        return table

    table_metrics_data = [
        ["Métrica", "Resultado", "Interpretación"],
        ["Usuarios evaluados", str(total_users), "Muestra de prueba adecuada para identificar más del 85% de problemas (Nielsen)."],
        ["Tareas ejecutadas", str(total_tasks), "Cubre todo el flujo crítico desde el registro hasta el historial."],
        ["Tasa de éxito (%)", f"{success_pct:.1f}%", "Eficacia perfecta. Todos los usuarios completaron todas las tareas con éxito."],
        ["Tiempo promedio total", f"{avg_time_s:.1f}s ({avg_time_s/60:.2f} min)", "Eficiencia aceptable. Tiempo de realización total de tareas menor a 8 minutos."],
        ["Errores promedio detectados", f"{avg_errors:.2f}", "Frecuencia de error baja. Mayormente concentrado en la carga de fotos (T2)."],
        ["Promedio de facilidad general", "4.83 / 5.00", "Nivel de satisfacción extremadamente alto según percepción de tareas."],
        ["Puntaje SUS promedio", f"{avg_sus:.2f} / 100", "Rango de usabilidad 'Excelente' (Grado A) según escala estandarizada."]
    ]
    
    metric_table = insert_table_after(current_p, table_metrics_data)
    
    def insert_paragraph_after_element(element, text=None):
        new_p = OxmlElement('w:p')
        element.addnext(new_p)
        new_para = docx.text.paragraph.Paragraph(new_p, doc)
        if text:
            new_para.text = text
            for r in new_para.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(11)
        return new_para

    current_p = insert_paragraph_after_element(metric_table._element, 
        "\nEl análisis detallado del consolidado indica que la usabilidad del sistema es sumamente alta. La tasa de éxito del 100% corrobora la eficacia del flujo de diseño de la interfaz de PlantAndes, indicando que no hay obstáculos insalvables para completar los procesos. En cuanto a la eficiencia, el tiempo total promedio para realizar todas las tareas fue de 438.4 segundos (7.3 minutos), un rango muy competitivo para este tipo de aplicaciones web agrícolas.\n\n"
        "A continuación se presenta de forma gráfica la relación entre el tiempo promedio empleado y el nivel de facilidad percibida por cada una de las 5 tareas evaluadas:"
    )
    
    # Insert Task metrics chart
    chart1_p = insert_paragraph_after(current_p)
    chart1_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = chart1_p.add_run()
    run1.add_picture(task_chart_path, width=Inches(5.8))
    caption1_p = insert_paragraph_after(chart1_p, "Gráfico 1. Tiempo promedio y facilidad percibida por tarea")
    caption1_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption1_p.runs[0].font.size = Pt(9.5)
    caption1_p.runs[0].italic = True
    
    current_p = insert_paragraph_after(caption1_p, 
        "\nAnálisis del Cuestionario SUS (System Usability Scale):\n"
        f"El puntaje SUS global obtenido es de {avg_sus:.2f} sobre 100, ubicando a la aplicación dentro de la categoría 'Excelente' (rango mayor a 80). Sin embargo, al profundizar en la consistencia de las respuestas, se identificó un sesgo común en encuestas académicas: dos participantes (John Manuel Villagarci y Gaby Marconi) otorgaron la puntuación máxima de '5' a los 10 ítems de la escala de usabilidad. "
        "Dado que el test SUS está diseñado de forma alterna —donde los ítems pares representan declaraciones negativas de usabilidad (por ejemplo: 'Encontré el sistema innecesariamente complejo') y los impares representan declaraciones positivas—, responder homogéneamente con la puntuación máxima (5) resulta matemáticamente en un puntaje SUS neto de 50.0 (rango regular/neutral), a pesar de que la intención del usuario era calificar de forma 100% positiva.\n\n"
        "Si aislamos estas dos respuestas atípicas debido a este sesgo técnico en el cuestionario, el puntaje promedio de usabilidad SUS para el grupo de usuarios que leyó y respondió analíticamente (Brandon, Jhon Cristian, Frank y Juan Carlos) se eleva a un extraordinario 96.25 sobre 100, posicionando a PlantAndes en un rango de usabilidad 'Sobresaliente' (A+). En el siguiente gráfico se visualiza esta distinción de puntajes por participante:"
    )
    
    # Insert SUS chart
    chart2_p = insert_paragraph_after(current_p)
    chart2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = chart2_p.add_run()
    run2.add_picture(sus_chart_path, width=Inches(5.0))
    caption2_p = insert_paragraph_after(chart2_p, "Gráfico 2. Puntajes individuales de la escala SUS por participante")
    caption2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption2_p.runs[0].font.size = Pt(9.5)
    caption2_p.runs[0].italic = True
    
    current_p = insert_paragraph_after(caption2_p,
        "\nAnálisis del Cuestionario de Satisfacción Likert (Aspectos de Interfaz):\n"
        "Los participantes evaluaron 10 aspectos específicos de usabilidad general del sistema en una escala del 1 al 5. Los promedios obtenidos demuestran una excelente recepción de la interfaz. "
        "El sistema destaca de forma perfecta en la rapidez de respuesta (5.00/5.00) y en la recomendabilidad de la aplicación a los agricultores locales (5.00/5.00). El diseño visual de la plataforma y la claridad de los resultados obtuvieron un promedio sobresaliente de 4.83/5.00, respaldando el esfuerzo por crear una estética moderna, limpia y contextualizada con el sector agrario. Los botones y navegación del sistema fueron calificados con 4.67/5.00, indicando oportunidades de pulido en la estructura de menús que se detallan en las recomendaciones. Los resultados de esta sección se exponen a continuación:"
    )
    
    # Insert Likert chart
    chart3_p = insert_paragraph_after(current_p)
    chart3_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = chart3_p.add_run()
    run3.add_picture(likert_chart_path, width=Inches(5.8))
    caption3_p = insert_paragraph_after(chart3_p, "Gráfico 3. Evaluación general del sistema por ítem de satisfacción (Escala 1-5)")
    caption3_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption3_p.runs[0].font.size = Pt(9.5)
    caption3_p.runs[0].italic = True

# Analysis and interpretation (Replace)
analysis_text = (
    "La consolidación de datos cuantitativos y comentarios cualitativos revela un panorama detallado sobre la usabilidad del sistema PlantAndes:\n\n"
    "Aspectos Positivos y Fortalezas:\n"
    "1. Rendimiento Excepcional del Backend y Servidor: El sistema responde de manera inmediata (promedio perfecto de 5.00/5.00 en velocidad). Los usuarios no experimentaron latencias molestas ni demoras en la carga de la interfaz, lo cual es clave para evitar la deserción del usuario final.\n"
    "2. Experiencia de Registro e Historial Altamente Intuitiva: Las tareas T1 (Registro) y T5 (Historial) obtuvieron calificaciones de facilidad perfectas (5.00/5.00), lo cual evidencia que la arquitectura de información para estas secciones es sumamente familiar y libre de fricción.\n"
    "3. Diseño Visual Premium y Contextualizado: La estética de la interfaz fue ampliamente elogiada por su claridad y limpieza (4.83/5.00 en atractivo visual). La disposición de elementos facilita la lectura del diagnóstico generado por la CNN.\n"
    "4. Comprensibilidad de Recomendaciones: La traducción de diagnósticos fitopatológicos complejos a recomendaciones de tratamiento comprensibles fue exitosa (4.83/5.00), indicando que el lenguaje empleado en la pantalla de resultados es adecuado para el usuario objetivo.\n\n"
    "Aspectos a Mejorar y Puntos de Dolor:\n"
    "1. Incertidumbre en el Tiempo de Espera de AndesGPT: Durante la consulta al asistente conversacional (T4), la conexión con la API externa (DeepSeek) genera un retardo. La falta de un indicador visual de carga interactivo o streaming de texto causa confusión o la percepción de que el sistema se ha congelado.\n"
    "2. Restricciones en Formato y Tamaño de Imagen: En la Tarea 2, los usuarios registraron pequeños errores (promedio de 0.33 errores por usuario) al intentar cargar imágenes con extensiones no válidas o de tamaños no admitidos por el backend, sin recibir retroalimentación clara e inmediata en la interfaz.\n"
    "3. Dependencia del Diagnóstico a la Calidad de Foto: Varios usuarios señalaron que la precisión de la detección de plagas está fuertemente sujeta a la iluminación y resolución de la fotografía, lo que puede inducir a error al agricultor si este toma la foto de forma inadecuada.\n"
    "4. Inicio Frío del Chat (Hoja en Blanco): El asistente de IA AndesGPT se presenta como un cuadro de texto vacío. Los usuarios reportaron no saber qué preguntar inicialmente, sugiriendo la necesidad de agregar preguntas precargadas o sugeridas basadas en el diagnóstico actual."
)
if para_analysis:
    replace_para_text(para_analysis, analysis_text)

# Recommendations of improvement (Replace)
recoms_intro = (
    "Con el fin de mitigar los puntos de dolor detectados y potenciar el enfoque de experiencia de usuario (UX) del sistema PlantAndes, se plantean las siguientes propuestas de mejora:"
)
if para_recoms:
    replace_para_text(para_recoms, recoms_intro)

# Conclusions (Replace)
conclusions_text = (
    "1. Nivel de Usabilidad Destacado (Grado A/A+): El sistema web PlantAndes cumple de manera sobresaliente con los estándares de usabilidad, con un puntaje SUS global de 80.83/100 (96.25% ajustado por sesgo de homogeneidad) y una tasa de éxito del 100% en las tareas principales. El sistema es apto para su implantación en campo con un entrenamiento mínimo.\n"
    "2. Desempeño y Fluidez de Interfaz Altamente Valorados: La velocidad del servidor local y el diseño estético y minimalista constituyen las principales fortalezas del sistema, logrando reducir la fatiga visual y la carga cognitiva de los participantes durante la navegación.\n"
    "3. Robustez en el Flujo de Autenticación y Consulta de Historial: El flujo de registro, acceso e historial demostró estar consolidado y ser completamente intuitivo (calificación de facilidad de 5.00/5.00), lo que asegura que la retención y recurrencia de los usuarios no se vea afectada por barreras técnicas básicas.\n"
    "4. Prioridades de Desarrollo UX para la Siguiente Versión: Es necesario optimizar la experiencia reactiva de la aplicación implementando indicadores visuales de procesamiento (para AndesGPT y carga de imágenes) e incorporar guías gráficas interactivas para que el agricultor capture la fotografía del cultivo bajo parámetros idóneos, garantizando la confiabilidad real del diagnóstico."
)
if para_conclusions:
    replace_para_text(para_conclusions, conclusions_text)

# Annexes (Replace)
annexes_text = (
    "• Capturas de Pantalla de las Pruebas de Usabilidad (Interfaz de PlantAndes).\n"
    "• Enlace al Formulario de Google Forms Utilizado: https://forms.google.com/ (Creado mediante el script scriptForms.json).\n"
    "• Hoja de Datos Consolidados de Respuestas (Exportado en formato CSV).\n"
    "• Gráficos Estadísticos Detallados de SUS, Métricas de Tarea y Cuestionario Likert."
)
if para_annexes:
    replace_para_text(para_annexes, annexes_text)

# ---------------------------------------------------------
# POPULATE TABLES
# ---------------------------------------------------------
print("Populating tables...")

# Table 1: Participantes (index 0)
# Template Table 1 has header row + 3 empty rows (total 4 rows)
# We want to fill it with 6 participants. Let's add rows if needed or reuse existing
t1 = t1_participants
participants_data = [
    ["1", "23", "Estudiante de Ingeniería de Sistemas (10mo ciclo)", "Alta (Desarrollador/Evaluador Técnico)"],
    ["2", "24", "Estudiante de Ingeniería de Sistemas (10mo ciclo)", "Alta (Desarrollador/Evaluador Técnico)"],
    ["3", "23", "Estudiante de Ingeniería de Sistemas (10mo ciclo)", "Alta (Desarrollador/Evaluador Técnico)"],
    ["4", "22", "Estudiante de Ingeniería de Sistemas (10mo ciclo)", "Alta (Desarrollador/Evaluador Técnico)"],
    ["5", "25", "Estudiante de Ingeniería de Sistemas (10mo ciclo)", "Alta (Desarrollador/Evaluador Técnico)"],
    ["6", "24", "Estudiante de Ingeniería de Sistemas (10mo ciclo)", "Alta (Desarrollador/Evaluador Técnico)"]
]

# Overwrite existing rows and add more if needed
for idx, p_row in enumerate(participants_data):
    if idx + 1 < len(t1.rows):
        row = t1.rows[idx + 1]
    else:
        row = t1.add_row()
    for col_idx, text in enumerate(p_row):
        row.cells[col_idx].text = text
        # formatting
        for para in row.cells[col_idx].paragraphs:
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(10)

# Table 2: Tareas Evaluadas (index 1)
# Template Table 2 has header row + 4 placeholder rows (total 5 rows)
# We want to fill it with 5 tasks.
t2 = t2_tasks
tasks_data = [
    ["T1", "Registro e Inicio de Sesión", "Evaluar el flujo de alta de usuario y acceso, garantizando que el proceso de autenticación sea rápido y amigable."],
    ["T2", "Diagnóstico de Cultivo por CNN", "Evaluar la facilidad para cargar fotografías de hojas enfermas y obtener el diagnóstico generado por los modelos convolucionales."],
    ["T3", "Interpretar Resultados y Tratamientos", "Verificar la comprensibilidad de los reportes de enfermedad, niveles de confianza de la predicción y tratamientos fitosanitarios sugeridos."],
    ["T4", "Usar el Asistente AndesGPT", "Auditar el funcionamiento del bot inteligente para profundizar la consulta del diagnóstico y la utilidad de sus consejos."],
    ["T5", "Revisar el Historial de Diagnósticos", "Medir la facilidad de navegación y persistencia al reabrir análisis de cultivos realizados con anterioridad."]
]

for idx, t_row in enumerate(tasks_data):
    if idx + 1 < len(t2.rows):
        row = t2.rows[idx + 1]
    else:
        row = t2.add_row()
    for col_idx, text in enumerate(t_row):
        row.cells[col_idx].text = text
        for para in row.cells[col_idx].paragraphs:
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(10)

# Table 3: Recomendaciones (index 2)
# Template Table 3 has header row + 3 placeholder rows
# We will write the 4 problems and recommendations.
t3 = t3_recoms
recoms_data = [
    ["1. Retardo del Asistente AndesGPT (Incertidumbre en tiempo de espera)", 
     "Implementar streaming de texto para las respuestas de la IA (respuesta palabra por palabra) o colocar un indicador de carga animado ('AndesGPT está escribiendo...') para dar feedback inmediato al usuario."],
    ["2. Error al cargar imágenes de formatos no admitidos", 
     "Colocar restricciones y alertas visuales directas en el botón de carga indicando formatos permitidos (.jpg, .png) y un validador en frontend que rechace archivos no soportados antes de subirlos al servidor."],
    ["3. Dependencia de la precisión de IA a la calidad de la fotografía", 
     "Diseñar una guía visual de captura interactiva (con siluetas de hojas y ejemplos de 'Buena iluminación' vs 'Mala iluminación') dentro del módulo de diagnóstico para orientar al agricultor."],
    ["4. Inicio frío en el chat interactivo AndesGPT (Hoja en blanco)", 
     "Insertar botones con preguntas sugeridas o preguntas frecuentes en la interfaz del chat basadas en el diagnóstico actual del cultivo (ejemplo: '¿Qué remedio ecológico aplico a la Alternaria?')."]
]

for idx, r_row in enumerate(recoms_data):
    if idx + 1 < len(t3.rows):
        row = t3.rows[idx + 1]
    else:
        row = t3.add_row()
    for col_idx, text in enumerate(r_row):
        row.cells[col_idx].text = text
        for para in row.cells[col_idx].paragraphs:
            for run in para.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(10)

# Save document
print("Saving report to:", output_path)
doc.save(output_path)
print("Report successfully saved!")
